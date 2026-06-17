# ============================================================
# 文档加载器 — 把各种格式的文件解析成纯文本
# 支持：PDF、DOCX、TXT、Markdown
# 输出：LoadedDocument 对象（包含按页的文本内容）
# ============================================================

from __future__ import annotations

from pathlib import Path
from typing import Optional

from loguru import logger  # 比 print 更好的日志工具，能看到时间戳和模块名


class DocumentPage:
    """代表文档中的一页。"""
    def __init__(self, text: str, page_number: int) -> None:
        self.text = text             # 这一页的纯文本内容
        self.page_number = page_number  # 页码（从1开始）


class LoadedDocument:
    """解析后的完整文档。"""
    def __init__(
        self,
        file_path: str,               # 源文件路径
        filename: str,                # 文件名
        pages: list[DocumentPage],    # 所有页的列表
        metadata: Optional[dict] = None,  # 额外的元信息
    ) -> None:
        self.file_path = file_path
        self.filename = filename
        self.pages = pages
        self.metadata = metadata or {}

    @property
    def full_text(self) -> str:
        """把所有页的文本拼成一段。"""
        return "\n".join(p.text for p in self.pages)

    @property
    def char_count(self) -> int:
        """总字符数。"""
        return sum(len(p.text) for p in self.pages)

    @property
    def page_count(self) -> int:
        """总页数。"""
        return len(self.pages)


def load_document(file_path: str) -> LoadedDocument:
    """根据文件后缀名自动选择加载器。"""
    path = Path(file_path)
    ext = path.suffix.lower()       # 获取扩展名，如 .pdf
    logger.info("Loading: {} ({})", path.name, ext)  # 打印日志

    if ext == ".pdf":
        return _load_pdf(file_path, path.name)
    elif ext == ".docx":
        return _load_docx(file_path, path.name)
    elif ext in (".txt", ".md"):
        return _load_text(file_path, path.name)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── PDF 解析 ────────────────────────────────────────────

def _load_pdf(file_path: str, filename: str) -> LoadedDocument:
    """使用 pypdf 库解析 PDF 文件，按页提取文本。"""
    try:
        import pypdf
    except ImportError:
        raise ImportError("pip install pypdf")  # 没装库就提示安装

    pages: list[DocumentPage] = []
    with open(file_path, "rb") as f:  # rb = 二进制读取
        reader = pypdf.PdfReader(f)    # 创建 PDF 读取器
        for i, page in enumerate(reader.pages):
            # 提取文本，如果提取不到就用空字符串
            text = (page.extract_text() or "").strip()
            if text:  # 只保留非空页
                pages.append(DocumentPage(text=text, page_number=i + 1))
    return LoadedDocument(file_path, filename, pages)


# ── DOCX 解析 ───────────────────────────────────────────

def _load_docx(file_path: str, filename: str) -> LoadedDocument:
    """使用 python-docx 库解析 Word 文档。"""
    try:
        import docx
    except ImportError:
        raise ImportError("pip install python-docx")

    doc = docx.Document(file_path)
    # 提取所有段落，去掉空行
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs) if paragraphs else ""
    # Word 没有"页"的概念，所以所有文本放在一页里
    pages = [DocumentPage(text=text, page_number=1)]
    return LoadedDocument(file_path, filename, pages)


# ── TXT / Markdown 解析 ────────────────────────────────

def _load_text(file_path: str, filename: str) -> LoadedDocument:
    """普通文本文件，直接全部读取。"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        # errors="replace" 遇到编码错误也不崩溃，替换成 �
        text = f.read()
    pages = [DocumentPage(text=text, page_number=1)]
    return LoadedDocument(file_path, filename, pages)
